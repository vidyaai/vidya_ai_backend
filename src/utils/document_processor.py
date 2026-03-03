import base64
import json
from typing import Any
from openai import OpenAI
from controllers.config import logger
import pypdf
import docx
from io import BytesIO
import csv
import html2text
import markdown
from pdf2image import convert_from_bytes


class DocumentProcessor:
    """Service for processing various document types and extracting text content"""

    def __init__(self, use_docling: bool = True):
        """
        Initialize document processor

        Args:
            use_docling: If True, use Docling for PDF extraction (faster, more reliable)
                        If False, use GPT-4o vision (legacy method)
        """
        self.client = OpenAI()
        self.model = "gpt-4o"  # Vision-capable model for PDF extraction
        self.use_docling = use_docling

        # Choose PDF extraction method
        pdf_extractor = self._extract_pdf_text_docling if use_docling else self._extract_pdf_text

        self.supported_types = {
            "application/pdf": pdf_extractor,
            "text/plain": self._extract_text,
            "application/msword": self._extract_doc_text,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_docx_text,
            "text/markdown": self._extract_markdown_text,
            "text/html": self._extract_html_text,
            "text/csv": self._extract_csv_text,
            "application/json": self._extract_json_text,
        }

    def extract_text_from_file(
        self, file_content: str, file_name: str, file_type: str
    ) -> str:
        """
        Extract text content from various file types

        Args:
            file_content: Base64 encoded file content
            file_name: Original file name
            file_type: MIME type or file extension

        Returns:
            Extracted text content
        """
        try:
            # Decode base64 content
            decoded_content = base64.b64decode(file_content)

            # Determine file type from extension if MIME type not recognized
            if file_type not in self.supported_types:
                extension = file_name.lower().split(".")[-1] if "." in file_name else ""
                file_type = self._get_mime_type_from_extension(extension)

            # Extract text based on file type
            if file_type in self.supported_types:
                extractor = self.supported_types[file_type]
                return extractor(decoded_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error extracting text from {file_name}: {str(e)}")
            raise

    def _get_mime_type_from_extension(self, extension: str) -> str:
        """Map file extensions to MIME types"""
        extension_map = {
            "pdf": "application/pdf",
            "txt": "text/plain",
            "doc": "application/msword",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "md": "text/markdown",
            "html": "text/html",
            "htm": "text/html",
            "csv": "text/csv",
            "json": "application/json",
        }
        return extension_map.get(extension, "text/plain")

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF files using Poppler and GPT-4o vision"""
        try:
            # Convert PDF pages to images using Poppler
            images = convert_from_bytes(content, dpi=200)
            num_pages = len(images)
            logger.info(f"Converted PDF to {num_pages} images")

            # Strategy 1: For PDFs <= 10 pages, extract all at once
            # Note: 10 pages balances speed vs reliability (5 was too slow, 15 caused refusals)
            if num_pages <= 10:
                return self._extract_pdf_chunk(images, 1, num_pages)

            # Strategy 2: For large PDFs, chunk into batches of 10 pages
            else:
                logger.info(f"Large PDF detected ({num_pages} pages). Using chunked extraction.")
                extracted_chunks = []
                chunk_size = 10  # Balanced for speed and reliability

                for start_idx in range(0, num_pages, chunk_size):
                    end_idx = min(start_idx + chunk_size, num_pages)
                    chunk_images = images[start_idx:end_idx]
                    start_page = start_idx + 1
                    end_page = end_idx

                    logger.info(f"Extracting pages {start_page}-{end_page}...")
                    chunk_text = self._extract_pdf_chunk(chunk_images, start_page, end_page)
                    extracted_chunks.append(chunk_text)

                logger.info(f"Successfully extracted all {num_pages} pages in {len(extracted_chunks)} chunks")
                return "\n\n".join(extracted_chunks)

        except Exception as e:
            logger.error(f"Error extracting PDF text with Poppler/GPT-4o: {str(e)}")
            # Fallback to PyPDF2 if vision extraction fails
            logger.info("Attempting fallback to PyPDF2 text extraction...")
            try:
                pdf_file = BytesIO(content)
                pdf_reader = pypdf.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            except Exception as fallback_error:
                logger.error(f"Fallback extraction also failed: {str(fallback_error)}")
                raise ValueError("Failed to extract text from PDF file")

    def _extract_pdf_chunk(self, images: list, start_page: int, end_page: int) -> str:
        """Extract text from a chunk of PDF pages using GPT-4o vision"""
        try:
            # Convert images to base64
            image_contents = []
            for image in images:
                buffered = BytesIO()
                image.save(buffered, format="PNG")
                img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
                image_contents.append(
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{img_base64}"},
                    }
                )

            # Build message content with text instruction followed by all images
            page_range = f"{start_page}-{end_page}" if start_page != end_page else f"{start_page}"
            num_images = len(images)
            user_content = [
                {
                    "type": "text",
                    "text": f"You are viewing {num_images} page image(s) from an academic document. Extract ALL text from these images.\n\nFor each image:\n1. Start with '--- Page N ---' where N is the page number (starting from {start_page})\n2. Extract ALL visible text exactly as it appears\n3. Maintain original formatting, structure, equations, and layout\n4. For diagrams, figures, charts, graphs, or circuits, add:\n   [DIAGRAM: Brief description including key labels and values]\n5. Preserve all mathematical equations, formulas, and technical terms\n6. Keep everything in sequential page order\n\nBegin extraction now:",
                }
            ]
            user_content.extend(image_contents)

            # API call to extract text from this chunk
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert OCR system that extracts text from document images. Your job is to read images and transcribe ALL visible text exactly as shown. Also describe any diagrams, charts, graphs, equations, or visual elements with technical precision. Never refuse - always extract the visible text from the provided images.",
                    },
                    {"role": "user", "content": user_content},
                ],
                max_tokens=16384,  # GPT-4o max completion tokens
            )

            extracted_text = response.choices[0].message.content

            # Check for refusal patterns in the response
            if extracted_text:
                refusal_patterns = [
                    "I'm sorry, I can't",
                    "I cannot",
                    "I'm unable to",
                    "I can't assist",
                    "I can't process",
                    "provide the pages individually",
                    "one at a time",
                    "smaller groups"
                ]

                extracted_lower = extracted_text.lower()
                for pattern in refusal_patterns:
                    if pattern.lower() in extracted_lower:
                        logger.error(f"GPT-4o refused to extract pages {page_range}. Response: {extracted_text[:200]}")
                        raise ValueError(f"GPT-4o refused extraction for pages {page_range}. Try reducing chunk size or checking page content.")

                logger.info(f"Successfully extracted pages {page_range} ({len(extracted_text)} chars)")
                return extracted_text.strip()
            else:
                raise ValueError(f"Empty response from GPT-4o for pages {page_range}")

        except Exception as e:
            logger.error(f"Error extracting pages {page_range}: {str(e)}")
            raise

    def _extract_pdf_text_docling(self, content: bytes) -> str:
        """
        Extract PDF text using Docling (fast, reliable, image detection)

        This is the new default method for PDF extraction.
        Benefits over GPT-4o vision:
        - 8.8x faster (35-41s vs 305s for 47-page PDF)
        - Free (no API costs)
        - More reliable (no refusals)
        - Detects all images with location markers
        - Better structure preservation (markdown output)

        Args:
            content: PDF file bytes

        Returns:
            Extracted markdown text with image markers
        """
        try:
            from utils.docling_processor import DoclingProcessor

            logger.info("Extracting PDF with Docling (fast mode)")

            # Use Docling without image descriptions (fast, free)
            # Image descriptions can be added later per-topic if needed
            processor = DoclingProcessor(enable_image_descriptions=False)
            markdown_text = processor.extract_text_from_pdf(
                content,
                describe_images=False
            )

            logger.info(
                f"Docling extraction complete: {len(markdown_text):,} chars, "
                f"{markdown_text.count('<!-- image -->')} images detected"
            )

            return markdown_text

        except Exception as e:
            logger.error(f"Docling extraction failed: {str(e)}")
            logger.info("Falling back to legacy GPT-4o vision extraction...")

            # Fallback to GPT-4o vision if docling fails
            return self._extract_pdf_text(content)

    def _extract_text(self, content: bytes) -> str:
        """Extract text from plain text files"""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file")

    def _extract_doc_text(self, content: bytes) -> str:
        """Extract text from DOC files (legacy Word format)"""
        # Note: python-docx doesn't support .doc files, only .docx
        # For .doc files, we'd need additional libraries like python-docx2txt or antiword
        # For now, we'll return an error message suggesting conversion
        raise ValueError(
            "Legacy .doc files are not supported. Please convert to .docx or save as plain text."
        )

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc_file = BytesIO(content)
            doc = docx.Document(doc_file)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise ValueError("Failed to extract text from DOCX file")

    def _extract_markdown_text(self, content: bytes) -> str:
        """Extract text from Markdown files"""
        try:
            md_text = content.decode("utf-8")
            # Convert markdown to HTML first, then to plain text
            html = markdown.markdown(md_text)
            h = html2text.HTML2Text()
            h.ignore_links = True
            return h.handle(html).strip()
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {str(e)}")
            return content.decode("utf-8", errors="ignore")

    def _extract_html_text(self, content: bytes) -> str:
        """Extract text from HTML files"""
        try:
            html_content = content.decode("utf-8")
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            return h.handle(html_content).strip()
        except Exception as e:
            logger.error(f"Error extracting HTML text: {str(e)}")
            return content.decode("utf-8", errors="ignore")

    def _extract_csv_text(self, content: bytes) -> str:
        """Extract text from CSV files"""
        try:
            csv_text = content.decode("utf-8")
            csv_file = BytesIO(csv_text.encode())
            reader = csv.reader(csv_text.splitlines())

            text = ""
            for row_num, row in enumerate(reader):
                if row_num == 0:
                    text += "Headers: " + ", ".join(row) + "\n\n"
                else:
                    text += "Row " + str(row_num) + ": " + ", ".join(row) + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting CSV text: {str(e)}")
            return content.decode("utf-8", errors="ignore")

    def _extract_json_text(self, content: bytes) -> str:
        """Extract text from JSON files"""
        try:
            json_text = content.decode("utf-8")
            json_data = json.loads(json_text)

            # Convert JSON to readable text format
            return self._json_to_text(json_data)
        except Exception as e:
            logger.error(f"Error extracting JSON text: {str(e)}")
            return content.decode("utf-8", errors="ignore")

    def _json_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert JSON data to readable text format"""
        text = ""
        prefix = "  " * indent

        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text += f"{prefix}{key}:\n"
                    text += self._json_to_text(value, indent + 1)
                else:
                    text += f"{prefix}{key}: {value}\n"
        elif isinstance(data, list):
            for i, item in enumerate(data):
                text += f"{prefix}Item {i + 1}:\n"
                text += self._json_to_text(item, indent + 1)
        else:
            text += f"{prefix}{data}\n"

        return text
